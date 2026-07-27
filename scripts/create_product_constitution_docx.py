from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Product_Constitution_Blinkr_Urun_Anayasasi.docx"


def set_run(run, *, bold=False, size=11, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, *, style=None, bold=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run(r, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run(r)
    return p


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Product Constitution")
    set_run(r, bold=True, size=22, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Blinkr Urun Anayasasi ve Karar Filtresi")
    set_run(r, bold=True, size=16, color="2E74B5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Kurucu urun yonetisimi belgesi")
    set_run(r, size=11, color="555555")


def build():
    doc = Document()
    configure_styles(doc)
    add_title_block(doc)

    doc.add_heading("Oncelik", level=1)
    add_para(
        doc,
        "Bu belge Blinkr'in en ust seviye urun anayasasini tanimlar. Tum ADR'ler, mimari kararlar, urun yol haritasi ve kapsam kararlari bu belgeyle uyumlu olmak zorundadir. Celiski durumunda bu belge onceliklidir.",
    )

    doc.add_heading("Baglam", level=1)
    add_para(
        doc,
        "Blinkr buyudukce mesajlasma, story, kisa video, sonsuz feed, gelismis profil ve genel sosyal etkilesim ozellikleri gundeme gelebilir. Bu ozellikler kontrolsuz eklenirse urun, harita merkezli yer karari platformu olmaktan cikarak genel sosyal medya uygulamasina yaklasir.",
    )

    doc.add_heading("Karar", level=1)
    add_para(
        doc,
        "Blinkr'in varlik amaci, insanlarin gercek dunyada bir yer hakkinda daha hizli, daha dogru ve daha guvenli karar vermesini saglamaktir.",
        bold=True,
    )
    add_para(
        doc,
        "Blinkr'in basarisi kullanicilarin uygulamada gecirdigi sureyle degil; yer kararlarinin kalitesi, guvenilirligi ve kullaniciya sagladigi zaman tasarrufuyla degerlendirilir.",
    )
    add_para(
        doc,
        "Yeni onerilen her ozellik MVP'ye veya urun yol haritasina alinmadan once yazili bir gerekceyle su sorular uzerinden degerlendirilir:",
    )
    for item in [
        "Kullanicinin bir yer hakkinda daha hizli karar vermesini sagliyor mu?",
        "Kararin dogrulugunu veya guvenilirligini artiriyor mu?",
        "Mahremiyeti ve kisisel guvenligi koruyor mu?",
        "Blinkr'in map-first ve place-first kimligini guclendiriyor mu?",
    ]:
        add_number(doc, item)
    add_para(
        doc,
        "Bu kriterlerle uyumu yazili olarak gosterilemeyen ozellikler MVP kapsamina alinmaz.",
        bold=True,
    )
    add_para(
        doc,
        "Bu belge Blinkr'in urun kapsami, mimari oncelikleri ve yol haritasi icin baglayici referans dokumandir.",
        bold=True,
    )

    doc.add_heading("Izleme Metrikleri", level=1)
    add_para(
        doc,
        "Asagidaki metrikler hedef degil, karar filtresinin dogru calisip calismadigini izlemek icin kullanilir:",
    )
    for item in [
        "Ortalama karar verme suresi.",
        "Taze sinyal kapsama orani.",
        "Guvenilir kaynak orani.",
        "Yanlis veya eski bilgi geri bildirimi.",
        "Guvenlik olayi orani.",
        "Kullanici zaman tasarrufu sinyali.",
        "Harita sonucundan gercek dunya eylemine gecis orani.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Istisna Kurali", level=1)
    add_para(
        doc,
        "Kullanicinin yer kararina dogrudan hizmet etmeyen bazi ozellikler; guvenlik, mevzuat uyumu, kotuye kullanim onleme, operasyonel surdurulebilirlik veya sistem guvenilirligi icin gerekli olabilir. Bu ozellikler kapsam icine alinabilir; ancak hangi riski azalttigi veya hangi zorunlulugu karsiladigi acikca belgelenmelidir.",
    )

    doc.add_heading("Mimari Etkiler", level=1)
    for item in [
        "Istemci deneyimi map-first kalir.",
        "Place modeli merkezi varlik olarak korunur.",
        "Icerikler mumkun oldugunca kisi koordinati yerine place, geo-cell veya yaklasik alan baglamina baglanir.",
        "Gateway/BFF bounds, place detail, freshness, source badge ve visibility policy sozlesmelerini one cikarir.",
        "Discovery/read model tazelik, guvenilirlik, mesafe ve gorunurluk politikalarini birlikte hesaplayacak sekilde tasarlanir.",
        "Event delivery, outbox/inbox ve projection guvenilirligi urun degerinin parcasidir; cunku gorunmeyen veya kaybolan sinyal yer kararini bozar.",
        "Mesajlasma, story, kisa video ve sonsuz genel feed ana mimariyi yonlendiren cekirdek capability kabul edilmez.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Negatif Sinirlar", level=1)
    add_para(doc, "Blinkr bilincli olarak sunlar olmayacaktir:")
    for item in [
        "Kullaniciyi uygulamada daha uzun tutmayi amaclayan genel sosyal medya urunu.",
        "Surekli kisi takip veya canli konum izleme uygulamasi.",
        "Isletme reklamlarini dogrulanmis yer sinyali gibi gizleyen pazarlama yuzeyi.",
        "Sonsuz feed ve eglence tuketimi merkezli icerik platformu.",
        "Guvenlik ve mahremiyet pahasina engagement buyuten sosyal ag.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Degisiklik Politikasi", level=1)
    add_para(
        doc,
        "Bu belge yalnizca urun vizyonunda bilincli ve stratejik bir degisiklik gerektiginde guncellenebilir.",
    )
    add_para(
        doc,
        "Her degisiklik; gerekcesi, beklenen etkileri, kullanici guvenligi ve mahremiyet sonuclari, urun kapsamina etkisi ve mevcut ADR'ler/mimari kararlarla uyumu belgelenerek yapilmalidir.",
    )

    doc.add_heading("Sonuclar", level=1)
    add_para(
        doc,
        "Tum urun, tasarim ve mimari kararlar bu belgeyle uyumlu olmalidir. Yeni bir ozellik 'hos olur mu?' diye degil, 'kullanicinin gercek dunyada daha iyi yer karari vermesini sagliyor mu?' diye degerlendirilir.",
    )

    doc.add_heading("Motto", level=1)
    add_para(
        doc,
        "Blinkr insanlarin uygulamada daha fazla zaman gecirmesi icin degil, gercek dunyada daha iyi yer kararlari vermesi icin vardir.",
        bold=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
